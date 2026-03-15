"""
Document parsing utilities.

Handles extraction of text content from various file formats:
    - PDF files (via pypdf)
    - DOCX files (via python-docx) - TODO
    - Plain text files - TODO
    - Markdown files - TODO

Used by the RAG pipeline for document ingestion and vectorization.
"""

import io
import structlog
from typing import Optional

from pypdf import PdfReader
from fastapi import UploadFile

from app.core.config import settings
from app.core.exceptions import DocumentProcessingException, BadRequestException

logger = structlog.get_logger(__name__)


async def _extract_text_from_pdf(file: UploadFile) -> str:
    """
    Extract text content from a PDF file.

    Args:
        file: FastAPI UploadFile object containing PDF data

    Returns:
        str: Extracted text content (pages joined by newlines)

    Raises:
        BadRequestException: If file is empty or unreadable
        DocumentProcessingException: If PDF parsing fails

    Notes:
        - File pointer is reset to beginning after reading
        - Empty pages are logged but don't cause failure
        - If no text is extractable, raises exception
    """
    log = logger.bind(filename=file.filename, content_type=file.content_type)
    log.info("pdf_extraction_started")

    try:
        # Read file content
        content = await file.read()
        if not content:
            raise BadRequestException(
                message="Uploaded file is empty", detail={"filename": file.filename}
            )

        # Validate file size
        file_size_mb = len(content) / (1024 * 1024)
        if file_size_mb > settings.MAX_UPLOAD_SIZE_MB:
            raise BadRequestException(
                message=f"File too large. Maximum size: {settings.MAX_UPLOAD_SIZE_MB} MB",
                detail={
                    "file_size_mb": file_size_mb,
                    "max_size_mb": settings.MAX_UPLOAD_SIZE_MB,
                },
            )

        # Parse PDF
        pdf = PdfReader(io.BytesIO(content))
        page_count = len(pdf.pages)

        log.info("pdf_parsed", page_count=page_count)

        # Extract text from each page
        text_pages = []
        empty_pages = []

        for page_num, page in enumerate(pdf.pages, start=1):
            try:
                extracted = page.extract_text()
                if extracted and extracted.strip():
                    text_pages.append(extracted)
                else:
                    empty_pages.append(page_num)
                    log.debug("pdf_page_empty", page=page_num)
            except Exception as e:
                log.warning("pdf_page_extraction_failed", page=page_num, error=str(e))
                continue

        # Reset file pointer for potential reuse
        await file.seek(0)

        # Validate extraction
        if not text_pages:
            raise DocumentProcessingException(
                message="Could not extract any readable text from PDF",
                detail={
                    "filename": file.filename,
                    "page_count": page_count,
                    "empty_pages": empty_pages,
                },
            )

        # Join pages with clear separator
        result = "\n\n".join(text_pages)

        log.info(
            "pdf_extraction_completed",
            page_count=page_count,
            pages_with_text=len(text_pages),
            empty_pages=len(empty_pages),
            text_length=len(result),
        )

        return result

    except (BadRequestException, DocumentProcessingException):
        # Re-raise our custom exceptions
        raise
    except Exception as e:
        log.error("pdf_extraction_error", error=str(e), error_type=type(e).__name__)
        raise DocumentProcessingException(
            message="Failed to parse PDF document",
            detail={"filename": file.filename, "error": str(e)},
        ) from e


async def _extract_text_from_docx(file: UploadFile) -> str:
    """
    Extract text content from a DOCX file.

    Args:
        file: FastAPI UploadFile object containing DOCX data

    Returns:
        str: Extracted text content

    Raises:
        DocumentProcessingException: If DOCX parsing fails
    """
    log = logger.bind(filename=file.filename)

    try:
        from docx import Document

        content = await file.read()
        if not content:
            raise BadRequestException("Uploaded file is empty")

        # Parse DOCX
        doc = Document(io.BytesIO(content))

        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)

        await file.seek(0)

        result = "\n\n".join(paragraphs + tables_text)

        if not result.strip():
            raise DocumentProcessingException("Could not extract any text from DOCX")

        log.info(
            "docx_extraction_completed",
            paragraph_count=len(paragraphs),
            table_count=len(doc.tables),
            text_length=len(result),
        )

        return result

    except ImportError:
        raise DocumentProcessingException(
            message="DOCX parsing not supported. Install python-docx.",
            detail={"install": "pip install python-docx"},
        )
    except Exception as e:
        log.error("docx_extraction_error", error=str(e))
        raise DocumentProcessingException(f"Failed to parse DOCX: {str(e)}") from e


async def _extract_text_from_txt(file: UploadFile) -> str:
    """
    Extract text from plain text file.

    Args:
        file: FastAPI UploadFile object containing text data

    Returns:
        str: File content as string
    """
    log = logger.bind(filename=file.filename)

    try:
        content = await file.read()
        if not content:
            raise BadRequestException("Uploaded file is empty")

        await file.seek(0)

        # Try UTF-8 first, fall back to latin-1
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            log.warning("utf8_decode_failed", trying="latin-1")
            text = content.decode("latin-1")

        log.info("txt_extraction_completed", text_length=len(text))
        return text

    except Exception as e:
        log.error("txt_extraction_error", error=str(e))
        raise DocumentProcessingException(f"Failed to read text file: {str(e)}") from e


async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text from uploaded file based on content type.

    Automatically detects file type and uses appropriate parser.

    Args:
        file: FastAPI UploadFile object

    Returns:
        str: Extracted text content

    Raises:
        BadRequestException: If file type is not supported
        DocumentProcessingException: If extraction fails
    """
    log = logger.bind(filename=file.filename, content_type=file.content_type)

    # Determine file type
    if file.content_type == "application/pdf" or file.filename.endswith(".pdf"):
        return await _extract_text_from_pdf(file)

    elif file.content_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ] or file.filename.endswith((".docx", ".doc")):
        return await _extract_text_from_docx(file)

    elif file.content_type == "text/plain" or file.filename.endswith(".txt"):
        return await _extract_text_from_txt(file)

    elif file.filename.endswith(".md"):
        return await _extract_text_from_txt(file)

    else:
        log.warning("unsupported_file_type")
        raise BadRequestException(
            message="Unsupported file type",
            detail={
                "content_type": file.content_type,
                "filename": file.filename,
                "supported_types": settings.ALLOWED_UPLOAD_EXTENSIONS,
            },
        )


def validate_file_extension(filename: str) -> bool:
    """
    Check if file extension is allowed.

    Args:
        filename: Name of file to validate

    Returns:
        bool: True if extension is allowed
    """
    if not filename:
        return False

    return any(
        filename.lower().endswith(ext) for ext in settings.ALLOWED_UPLOAD_EXTENSIONS
    )


def get_file_extension(filename: str) -> Optional[str]:
    """
    Extract file extension from filename.

    Args:
        filename: Name of file

    Returns:
        str: File extension (with dot), or None if no extension
    """
    if not filename or "." not in filename:
        return None
    return "." + filename.rsplit(".", 1)[-1].lower()
